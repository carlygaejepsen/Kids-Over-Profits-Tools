import pandas as pd
from docx import Document

# Load CSV
df = pd.read_csv("utah_citations_8-24-2025.csv")

doc = Document()

for _, row in df.iterrows():
    # Find last non-empty inspection date
    dates = [row[i] for i in df.columns if "Inspection" in i and "Date" in i]
    last_date = max([d for d in dates if pd.notna(d) and d != "None"], default="None")

    # Match findings for that date
    findings = "None"
    for i in range(0, len(df.columns)):
        if df.columns[i].endswith("Date") and row[i] == last_date:
            findings = row[i+2] if pd.notna(row[i+2]) else "None"

    doc.add_paragraph(f"Name: {row['Name']}")
    doc.add_paragraph(f"Address: {row['Address']}")
    doc.add_paragraph(f"Date of last inspection: {last_date}")
    doc.add_paragraph(f"Results: {findings}")
    doc.add_paragraph("")  # blank line between entries

doc.save("utah_citations_report.docx")